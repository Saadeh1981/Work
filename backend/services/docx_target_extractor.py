from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document

from backend.services.evidence_builder import build_evidence
from backend.services.confidence_builder import score_text_block

def _clean_text(value: str | None) -> str:
    if not value:
        return ""
    return " ".join(value.replace("\xa0", " ").split()).strip()


def extract_docx_content(path: str | Path) -> dict[str, Any]:
    try:
        doc = Document(str(path))
    except Exception as exc:
        return {
            "status": "error",
            "paragraphs": [],
            "tables": [],
            "paragraph_evidence": [],
            "table_evidence": [],
            "paragraph_count": 0,
            "table_count": 0,
            "row_count": 0,
            "error": str(exc),
        }

    file_name = Path(path).name
    paragraphs: list[str] = []
    tables: list[list[list[str]]] = []
    paragraph_evidence: list[dict[str, Any]] = []
    table_evidence: list[dict[str, Any]] = []

    for i, para in enumerate(doc.paragraphs):
        text = _clean_text(para.text)
        if text:
            paragraphs.append(text)
            paragraph_evidence.append(
                build_evidence(
                    file_name=file_name,
                    file_type="docx",
                    snippet=text,
                    method="docx_paragraph",
                    block_index=i,
                    page=None,
                    confidence=score_text_block(
                        text=text,
                        method="docx_paragraph",
                        file_type="docx",
                    ),
                )
            )

    for t_index, table in enumerate(doc.tables):
        table_rows: list[list[str]] = []

        for r_index, row in enumerate(table.rows):
            cells = [_clean_text(cell.text) for cell in row.cells]
            if any(cells):
                table_rows.append(cells)

                row_text = " | ".join(cells)

                table_evidence.append(
                    build_evidence(
                        file_name=file_name,
                        file_type="docx",
                        snippet=row_text,
                        method="docx_table_row",
                        block_index=r_index,
                        page=None,
                        confidence=score_text_block(
                            text=row_text,
                            method="docx_table_row",
                            file_type="docx",
                        ),
                    )
                )

        if table_rows:
            tables.append(table_rows)

    row_count = sum(len(table) for table in tables)

    return {
        "status": "ok",
        "paragraphs": paragraphs,
        "tables": tables,
        "paragraph_evidence": paragraph_evidence,
        "table_evidence": table_evidence,
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "row_count": row_count,
        "error": None,
    }