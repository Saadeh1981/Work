from __future__ import annotations

from typing import List

import docx

from backend.schemas.summary import FileSummary, FileSection
from backend.services.preflight_service import DocProfile
from backend.services.summary_classifier import classify_text


MAX_PARAGRAPHS = 60
MAX_TABLES = 10
MAX_TABLE_ROWS = 15
MAX_TABLE_COLS = 12


def summarize_docx(prof: DocProfile) -> FileSummary:
    sections: List[FileSection] = []
    domains = set()
    content_types = set()
    entities = set()
    attributes = set()
    relationships = set()
    warnings = list(prof.warnings or [])

    try:
        doc = docx.Document(prof.path)
    except Exception as e:
        warnings.append(f"DOCX open failed in summary: {type(e).__name__}")
        return FileSummary(
            file_name=prof.name,
            file_type="docx",
            path=prof.path,
            text_layer_present=True,
            scan_likelihood=0.0,
            warnings=warnings,
            recommended_actions=[prof.plan or "native_parse"],
        )

    paragraph_sections = summarize_docx_paragraphs(doc)
    table_sections = summarize_docx_tables(doc)

    sections.extend(paragraph_sections)
    sections.extend(table_sections)

    for sec in sections:
        domains.update(sec.domains)
        content_types.update(sec.content_types)
        entities.update(sec.entities)
        attributes.update(sec.attributes)
        relationships.update(sec.relationships)

    return FileSummary(
        file_name=prof.name,
        file_type="docx",
        path=prof.path,
        text_layer_present=True,
        scan_likelihood=0.0,
        domains=sorted(domains),
        content_types=sorted(content_types),
        entities=sorted(entities),
        attributes=sorted(attributes),
        relationships=sorted(relationships),
        sections=sections,
        recommended_actions=["inspect_paragraphs_and_tables", "extract_detected_sections"],
        warnings=warnings,
    )


def summarize_docx_paragraphs(doc) -> List[FileSection]:
    sections: List[FileSection] = []

    paragraph_texts: List[str] = []
    heading_texts: List[str] = []

    for para in doc.paragraphs[:MAX_PARAGRAPHS]:
        text = (para.text or "").strip()
        if not text:
            continue

        style_name = ""
        try:
            style_name = (para.style.name or "").lower()
        except Exception:
            style_name = ""

        if "heading" in style_name or "title" in style_name:
            heading_texts.append(text)

        paragraph_texts.append(text)

    if not paragraph_texts and not heading_texts:
        return sections

    combined_text = "\n".join(heading_texts + paragraph_texts)
    result = classify_text(combined_text)

    sections.append(
        FileSection(
            label=pick_docx_label(result, fallback="docx_paragraphs"),
            confidence=estimate_docx_confidence(result),
            signals=build_docx_signals("paragraphs", result, heading_texts),
            domains=result.get("domains", []),
            content_types=result.get("content_types", []),
            entities=result.get("entities", []),
            attributes=result.get("attributes", []),
            relationships=result.get("relationships", []),
        )
    )

    return sections


def summarize_docx_tables(doc) -> List[FileSection]:
    sections: List[FileSection] = []

    for table_index, table in enumerate(doc.tables[:MAX_TABLES], start=1):
        sample_text = build_table_sample_text(table)
        if not sample_text.strip():
            continue

        result = classify_text(sample_text)

        sections.append(
            FileSection(
                label=pick_docx_label(result, fallback=f"docx_table_{table_index}"),
                confidence=estimate_docx_confidence(result),
                signals=build_docx_signals(f"table_{table_index}", result, []),
                domains=result.get("domains", []),
                content_types=result.get("content_types", []),
                entities=result.get("entities", []),
                attributes=result.get("attributes", []),
                relationships=result.get("relationships", []),
            )
        )

    return sections


def build_table_sample_text(table) -> str:
    lines: List[str] = []

    for row_idx, row in enumerate(table.rows[:MAX_TABLE_ROWS], start=1):
        values: List[str] = []

        for cell in row.cells[:MAX_TABLE_COLS]:
            text = (cell.text or "").strip()
            if text:
                values.append(text)

        if values:
            lines.append(" | ".join(values))

    return "\n".join(lines)


def pick_docx_label(result: dict, fallback: str) -> str:
    content_types = result.get("content_types", [])
    entities = result.get("entities", [])

    if content_types:
        return content_types[0]
    if entities:
        return entities[0]
    return fallback


def estimate_docx_confidence(result: dict) -> float:
    score_groups = result.get("scores", {})

    total_hits = 0
    for group_scores in score_groups.values():
        total_hits += sum(group_scores.values())

    if total_hits >= 8:
        return 0.9
    if total_hits >= 5:
        return 0.8
    if total_hits >= 3:
        return 0.7
    if total_hits >= 1:
        return 0.6
    return 0.4


def build_docx_signals(section_name: str, result: dict, headings: List[str]) -> List[str]:
    signals = [f"section:{section_name}"]

    for heading in headings[:3]:
        signals.append(f"heading:{heading[:80]}")

    for domain in result.get("domains", [])[:3]:
        signals.append(f"domain:{domain}")

    for content_type in result.get("content_types", [])[:3]:
        signals.append(f"content_type:{content_type}")

    for entity in result.get("entities", [])[:5]:
        signals.append(f"entity:{entity}")

    return signals