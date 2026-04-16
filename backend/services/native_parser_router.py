# backend/services/native_parser_router.py
from __future__ import annotations

from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Any, Optional
from io import BytesIO

import fitz  # PyMuPDF

try:
    import docx  # python-docx
except Exception:
    docx = None

try:
    import openpyxl
except Exception:
    openpyxl = None


@dataclass
class TextBlock:
    text: str
    bbox: Optional[list[float]] = None
    block_type: str = "text"


@dataclass
class TableBlock:
    rows: list[list[Any]]
    bbox: Optional[list[float]] = None


@dataclass
class NativePage:
    page_or_sheet: str
    raw_text: str
    text_blocks: list[TextBlock]
    tables: list[TableBlock]


@dataclass
class NativeParseResult:
    file_type: str
    method: str
    native_text_detected: bool
    page_count: Optional[int]
    content: list[NativePage]


def parse_native(
    filename: str,
    file_bytes: bytes,
    max_pages: int | None = 25,
    include_blocks: bool = False,
) -> dict:
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        result = _parse_pdf_native(
            file_bytes=file_bytes,
            max_pages=max_pages,
            include_blocks=include_blocks,
            sample_chars_per_page=4000,
        )
    elif ext == ".docx":
        result = _parse_docx_native(file_bytes)
    elif ext == ".xlsx":
        result = _parse_xlsx_native(file_bytes)
    else:
        result = NativeParseResult(
            file_type="other",
            method="native",
            native_text_detected=False,
            page_count=None,
            content=[],
        )

    return asdict(result)


def _parse_pdf_native(
    file_bytes: bytes,
    max_pages: int | None = 25,
    include_blocks: bool = False,
    sample_chars_per_page: int = 4000,
    max_total_chars: int = 200_000,
    max_blocks_per_page: int = 150,
    max_block_chars: int = 800,
) -> NativeParseResult:
    doc = fitz.open(stream=file_bytes, filetype="pdf")

    total_pages = doc.page_count
    n = total_pages if max_pages is None else min(total_pages, max_pages)

    pages: list[NativePage] = []
    total_chars = 0

    for i in range(n):
        page = doc.load_page(i)

        # Text layer extraction, cheap
        raw_text_full = (page.get_text("text") or "").strip()
        total_chars += len(raw_text_full)

        # Keep response payload bounded
        raw_text = raw_text_full[:sample_chars_per_page]

        # Optional blocks, bounded
        blocks: list[TextBlock] = []
        if include_blocks:
            try:
                raw_blocks = page.get_text("blocks") or []
            except Exception:
                raw_blocks = []

            kept = 0
            for b in raw_blocks:
                if kept >= max_blocks_per_page:
                    break

                # (x0, y0, x1, y1, text, block_no, block_type)
                x0, y0, x1, y1 = b[0], b[1], b[2], b[3]
                text = b[4] if len(b) > 4 else ""
                t = (text or "").strip()
                if not t:
                    continue

                blocks.append(
                    TextBlock(
                        text=t[:max_block_chars],
                        bbox=[float(x0), float(y0), float(x1), float(y1)],
                    )
                )
                kept += 1

        # Phase A: no table detection
        tables: list[TableBlock] = []

        pages.append(
            NativePage(
                page_or_sheet=str(i + 1),
                raw_text=raw_text,
                text_blocks=blocks,
                tables=tables,
            )
        )

        # Global cap to prevent runaway PDFs
        if total_chars >= max_total_chars:
            break

    doc.close()

    return NativeParseResult(
        file_type="pdf",
        method="native",
        native_text_detected=(total_chars > 200),
        page_count=len(pages),
        content=pages,
    )


    doc.close()

    native_text_detected = total_chars > 200
    return NativeParseResult(
        file_type="pdf",
        method="native",
        native_text_detected=native_text_detected,
        page_count=n,
        content=pages,
    )



def _parse_docx_native(file_bytes: bytes) -> NativeParseResult:
    if docx is None:
        return NativeParseResult("docx", "native", False, None, [])

    d = docx.Document(BytesIO(file_bytes))

    text_parts: list[str] = []
    blocks: list[TextBlock] = []

    for p in d.paragraphs:
        t = (p.text or "").strip()
        if not t:
            continue
        text_parts.append(t)
        blocks.append(TextBlock(text=t))

    tables: list[TableBlock] = []
    for tb in d.tables:
        rows: list[list[Any]] = []
        for r in tb.rows:
            rows.append([(c.text or "").strip() for c in r.cells])
        if rows:
            tables.append(TableBlock(rows=rows))

    raw_text = "\n".join(text_parts).strip()

    return NativeParseResult(
        file_type="docx",
        method="native",
        native_text_detected=bool(raw_text) or bool(tables),
        page_count=1,
        content=[NativePage(page_or_sheet="1", raw_text=raw_text, text_blocks=blocks, tables=tables)],
    )


def _parse_xlsx_native(
    file_bytes: bytes,
    max_rows_per_sheet: int = 500,
    max_cells_per_sheet: int = 50_000,
    max_total_chars: int = 200_000,
) -> NativeParseResult:
    if openpyxl is None:
        return NativeParseResult("xlsx", "native", False, None, [])

    try:
        wb = openpyxl.load_workbook(BytesIO(file_bytes), data_only=True, read_only=True)
    except Exception as e:
        return NativeParseResult(
            file_type="xlsx",
            method="native",
            native_text_detected=False,
            page_count=0,
            content=[
                NativePage(
                    page_or_sheet="error",
                    raw_text=f"Failed to open workbook: {type(e).__name__}: {e}",
                    text_blocks=[],
                    tables=[],
                )
            ],
        )

    content: list[NativePage] = []
    total_chars = 0

    for sh in wb.worksheets:
        rows_out: list[list[Any]] = []
        text_parts: list[str] = []
        cells_seen = 0
        rows_seen = 0

        for row in sh.iter_rows(values_only=True):
            if rows_seen >= max_rows_per_sheet:
                break
            row_list = list(row)

            if all(v is None or (isinstance(v, str) and not v.strip()) for v in row_list):
                continue

            rows_out.append(row_list)
            rows_seen += 1

            for v in row_list:
                if v is None:
                    continue
                s = v.strip() if isinstance(v, str) else str(v)
                if not s:
                    continue

                text_parts.append(s)
                total_chars += len(s)
                cells_seen += 1

                if cells_seen >= max_cells_per_sheet or total_chars >= max_total_chars:
                    break

            if cells_seen >= max_cells_per_sheet or total_chars >= max_total_chars:
                break

        raw_text = "\n".join(text_parts).strip()
        tables = [TableBlock(rows=rows_out)] if rows_out else []

        content.append(
            NativePage(
                page_or_sheet=f"Sheet: {sh.title}",
                raw_text=raw_text,
                text_blocks=[],
                tables=tables,
            )
        )

        if total_chars >= max_total_chars:
            break

    try:
        wb.close()
    except Exception:
        pass

    return NativeParseResult(
        file_type="xlsx",
        method="native",
        native_text_detected=any(p.raw_text or p.tables for p in content),
        page_count=len(content),
        content=content,
    )
def build_preflight_summary(native_result: dict, max_sample_chars: int = 12_000) -> dict:
    file_type = native_result.get("file_type", "other")
    native_text_detected = bool(native_result.get("native_text_detected", False))
    page_count = native_result.get("page_count", None)

    content = native_result.get("content", []) or []

    # sample up to first 3 pages/sheets
    sample_parts: list[str] = []
    sampled_units = 0
    for item in content[:3]:
        txt = (item.get("raw_text") or "").strip()
        if not txt:
            continue
        sample_parts.append(txt)
        sampled_units += 1

        if sum(len(x) for x in sample_parts) >= max_sample_chars:
            break

    sample_text = "\n\n".join(sample_parts)[:max_sample_chars].strip()

    notes: list[str] = []
    if file_type == "pdf" and not native_text_detected:
        notes.append("No native text detected, likely scanned, plan OCR.")
    if file_type == "xlsx":
        notes.append("page_count represents sheets.")
    if file_type == "docx":
        notes.append("page_count is logical document page = 1.")

    return {
        "file_type": file_type,
        "method": "native_preflight",
        "page_count": page_count,
        "native_text_detected": native_text_detected,
        "sampled_units": sampled_units,
        "sample_text": sample_text,
        "notes": notes,
    }
